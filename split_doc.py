import json
from google import genai
from dotenv import load_dotenv
from docx import Document

load_dotenv()

client = genai.Client()



def ask_gemini(prompt):
    print(f"split_doc.py: {prompt[:100]}")
    response = client.models.generate_content(
        model="gemini-3-flash-preview",
        contents=prompt,
    )
    return response.text


def extract_json(text):
    text = text.strip()

    if text.startswith("```json"):
        text = text[7:]
    if text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]

    text = text.strip()

    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1:
        text = text[start:end + 1]

    try:
        data = json.loads(text)
        return data
    except Exception as e:
        print(f"Error parsing JSON: {e}")
        print(f"Original text: {text}")
        raise

def split_sections_with_llm(full_text):
    prompt = f"""
You are given the raw text of a document.

Your task:
Split the document into LARGE logical sections.

Rules:
- Use the table of contents if it exists.
- The first section should usually include the cover/title page and table of contents.
- Sections should be large chapter-level sections, not tiny sub-sections.
- Preserve the original text exactly as much as possible.
- Do not summarize.
- Do not rewrite.
- Do not omit content.
- Every part of the document must belong to exactly one section.

Return valid JSON only in this format:
{{
  "sections": [
    {{
      "section_title": "Document Header and Table of Contents",
      "text": "full original text of that section"
    }},
    {{
      "section_title": "כללי",
      "text": "full original text of that section"
    }}
  ]
}}

Document text:
{full_text}
"""
    text = ask_gemini(prompt)
    data = extract_json(text)
    return data["sections"]


def read_docx(path):
    print(f"Reading document from {path}...")
    doc = Document(path)
    blocks = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            blocks.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                blocks.append(" | ".join(cells))

    return blocks

def split_doc_to_sections(doc_path):
    print("loading document...")
    blocks = read_docx(doc_path)
    full_text = "\n\n".join(blocks)

    print(f"full document chars: {len(full_text)}")

    print("splitting document with llm...")
    sections = split_sections_with_llm(full_text)
    return sections

if __name__ == "__main__":
    doc_path = "uploads/c.docx"
    sections = split_doc_to_sections(doc_path)
    print(json.dumps(sections, ensure_ascii=False, indent=2))